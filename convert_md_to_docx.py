#!/usr/bin/env python3
"""
Convert Markdown file to Word document using python-docx
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def markdown_to_docx(md_file, docx_file):
    doc = Document()
    
    # Title
    doc.add_heading('AWS CRT Standalone Benchmark Analysis', 0)
    
    with open(md_file, 'r') as f:
        content = f.read()
    
    # Split by lines
    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], 0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], 1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], 2)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], 3)
        elif line.startswith('##### '):
            doc.add_heading(line[6:], 4)
        elif line.startswith('###### '):
            doc.add_heading(line[7:], 5)
        
        # Horizontal rule
        elif line.strip() == '---':
            doc.add_paragraph('_' * 50)
        
        # Empty line
        elif not line.strip():
            doc.add_paragraph()
        
        # Lists
        elif line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\. ', line):
            p = doc.add_paragraph(line[3:], style='List Number')
        
        # Bold
        elif '**' in line:
            # Handle bold text
            parts = re.split(r'\*\*(.*?)\*\*', line)
            p = doc.add_paragraph()
            for j, part in enumerate(parts):
                if j % 2 == 1:  # Bold text
                    run = p.add_run(part)
                    run.bold = True
                else:
                    p.add_run(part)
        
        # Code blocks
        elif line.startswith('```'):
            # Skip code block markers
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            code = '\n'.join(code_lines)
            p = doc.add_paragraph(code)
            run = p.runs[0]
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        
        # Regular paragraph
        elif line.strip():
            doc.add_paragraph(line)
        
        i += 1
    
    doc.save(docx_file)
    print(f"Converted {md_file} to {docx_file}")

if __name__ == '__main__':
    import sys
    if len(sys.argv) != 3:
        print("Usage: python convert_md_to_docx.py <input.md> <output.docx>")
        sys.exit(1)
    
    markdown_to_docx(sys.argv[1], sys.argv[2])
